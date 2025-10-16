import os
import pandas as pd
import discord
from discord.ext import commands
from utils import db


class DbAdmin(commands.Cog):
    def __init__(self, bot):
        self.bot = bot

    # ========================
    # 🔹 Group Command
    # ========================
    @commands.group(name="db", invoke_without_command=True)
    async def db_group(self, ctx):
        await ctx.send(
            "🧩 **Database Commands**\n"
            "• `!db checkschema` → cek semua tabel & kolom\n"
            "• `!db resettable <nama>` → hapus tabel tertentu\n"
            "• `!db exportitems <csv/xlsx/docx/json>` → ekspor item database\n"
            "• `!db export <nama_tabel> <csv/xlsx/docx/json>` → ekspor tabel tertentu\n"
            "• `!db exportall <csv/xlsx>` → ekspor semua tabel utama"
        )

    # ========================
    # 📘 Cek Schema Database
    # ========================
    @db_group.command(name="checkschema")
    async def checkschema(self, ctx):
        """Cek semua tabel & kolom lalu kirim hasilnya sebagai file .txt"""
        guild_id = ctx.guild.id
        schema = db.check_schema(guild_id)

        if not schema:
            return await ctx.send("❌ Tidak ada tabel yang ditemukan di database ini.")

        # Siapkan isi file
        lines = [f"📘 Database Schema – Guild ID: {guild_id}\n"]
        for table, cols in schema.items():
            lines.append(f"\n=== {table.upper()} ===")
            for col in cols:
                lines.append(f"• {col}")
        text = "\n".join(lines)

        # Simpan ke file sementara
        os.makedirs("/tmp", exist_ok=True)
        filename = f"/tmp/schema_{guild_id}.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)

        await ctx.send(
            content=f"✅ **Schema database guild {ctx.guild.name}** berhasil diekspor.",
            file=discord.File(filename)
        )

    @db_group.command(name="usebackup")
    @commands.has_permissions(administrator=True)
    async def usebackup(self, ctx):
        """🔁 Gunakan DB lama (../data/) untuk guild ini."""
        import shutil, os
        guild_id = ctx.guild.id
        old_path = f"../data/narator_{guild_id}.db"
        new_path = f"./data/narator_{guild_id}.db"

        if not os.path.exists(old_path):
            return await ctx.send(f"❌ Tidak ditemukan file lama: {old_path}")

        os.makedirs("./data", exist_ok=True)
        shutil.copy2(old_path, new_path)
        size = os.path.getsize(new_path) / 1024
        await ctx.send(f"✅ DB lama berhasil disalin ke lokasi aktif.\n📦 `{new_path}` ({size:.1f} KB)")

    # ========================
    # 🔍 Cek Info DB Aktif (baru)
    # ========================
    @db_group.command(name="info")
    @commands.has_permissions(administrator=True)
    async def db_info(self, ctx):
        """🔍 Lihat path file database aktif & jumlah data utama."""
        from utils.db import get_db_path, fetchall

        guild_id = ctx.guild.id
        path = get_db_path(guild_id)

        # Hitung isi tabel penting
        def count_rows(table):
            try:
                return len(fetchall(guild_id, f"SELECT * FROM {table}"))
            except Exception:
                return 0

        data = {
            "characters": count_rows("characters"),
            "npcs": count_rows("npc"),
            "items": count_rows("items"),
            "quests": count_rows("quests"),
            "factions": count_rows("factions"),
            "favors": count_rows("favors"),
            "hollow_nodes": count_rows("hollow_nodes"),
            "hollow_events": count_rows("hollow_events"),
            "hollow_visitors": count_rows("hollow_visitors"),
        }

        embed = discord.Embed(
            title="🧩 Database Info",
            description=f"📘 Guild ID: `{guild_id}`",
            color=discord.Color.blurple()
        )
        embed.add_field(name="📁 DB Path", value=f"`{path}`", inline=False)
        for k, v in data.items():
            embed.add_field(name=k.replace("_", " ").title(), value=str(v), inline=True)
        embed.set_footer(text="Technonesia System — Database Inspector")

        await ctx.send(embed=embed)
    
        # ========================
    # 📂 List Semua File Database (baru)
    # ========================
    @db_group.command(name="listfiles")
    @commands.has_permissions(administrator=True)
    async def db_listfiles(self, ctx):
        """📂 Lihat semua file database (.db) yang tersimpan di folder /data"""
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data"))
        if not os.path.exists(base_dir):
            return await ctx.send("❌ Folder `/data` belum dibuat.")

        db_files = [f for f in os.listdir(base_dir) if f.endswith(".db")]
        if not db_files:
            return await ctx.send("📭 Tidak ada file database ditemukan di `/data`.")

        embed = discord.Embed(
            title="📂 Database Files in /data",
            color=discord.Color.green()
        )

        for f in db_files:
            path = os.path.join(base_dir, f)
            size_kb = os.path.getsize(path) / 1024
            mtime = os.path.getmtime(path)
            from datetime import datetime
            modified = datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")

            embed.add_field(
                name=f"📘 {f}",
                value=f"🗜 {size_kb:.1f} KB\n🕒 {modified}",
                inline=False
            )

        embed.set_footer(text="Technonesia System — DB File Viewer")
        await ctx.send(embed=embed)
    
    # ========================
    # ⚙️ Inisialisasi Database
    # ========================
    @db_group.command(name="init")
    @commands.has_permissions(administrator=True)
    async def init_database(self, ctx):
        """Inisialisasi struktur database guild (buat tabel Hollow, dsb)."""
        guild_id = ctx.guild.id
        try:
            db.init_db(guild_id)
            await ctx.send(f"✅ Database untuk guild **{ctx.guild.name}** sudah diinisialisasi / disinkronkan.")
        except Exception as e:
            await ctx.send(f"❌ Gagal inisialisasi database: {e}")

    # ========================
    # 🗑️ Reset Tabel
    # ========================
    @db_group.command(name="resettable")
    @commands.has_permissions(administrator=True)
    async def resettable(self, ctx, *, table: str):
        """Hapus satu tabel lalu kosongkan"""
        guild_id = ctx.guild.id
        try:
            db.execute(guild_id, f"DROP TABLE IF EXISTS {table}")
            await ctx.send(f"✅ Tabel `{table}` sudah dihapus.")
        except Exception as e:
            await ctx.send(f"❌ Gagal drop tabel `{table}`: {e}")

    # ========================
    # 📤 Export Item Data (lama)
    # ========================
    @db_group.command(name="exportitems")
    @commands.has_permissions(administrator=True)
    async def exportitems(self, ctx, fmt: str = "csv"):
        """📦 Export semua item ke file (csv/xlsx/docx/json)."""
        guild_id = ctx.guild.id
        rows = db.fetchall(guild_id, "SELECT * FROM items")

        if not rows:
            return await ctx.send("❌ Tidak ada data item di database!")

        df = pd.DataFrame(rows)
        os.makedirs("/tmp", exist_ok=True)
        filename = f"/tmp/items_{guild_id}.{fmt.lower()}"

        try:
            fmt = fmt.lower()

            if fmt == "csv":
                df.to_csv(filename, index=False)
            elif fmt == "xlsx":
                df.to_excel(filename, index=False)
            elif fmt == "json":
                df.to_json(filename, orient="records", indent=2)
            elif fmt == "docx":
                from docx import Document
                doc = Document()
                doc.add_heading("Technonesia Item Catalog", level=1)

                for _, r in df.iterrows():
                    name = r.get("name", "?")
                    rarity = r.get("rarity", "?")
                    desc = r.get("desc", "")
                    doc.add_heading(f"{name} ({rarity})", level=2)
                    doc.add_paragraph(
                        f"Type: {r.get('type', '?')} | "
                        f"Value: {r.get('value', '?')} | "
                        f"Weight: {r.get('weight', '?')}"
                    )
                    if desc:
                        doc.add_paragraph(desc)
                    doc.add_paragraph("—")

                doc.save(filename)
            else:
                return await ctx.send("❌ Format tidak dikenal! Gunakan: csv / xlsx / docx / json.")

            await ctx.send(
                content=f"✅ Data item diekspor sebagai **{fmt.upper()}**",
                file=discord.File(filename),
            )

        except Exception as e:
            await ctx.send(f"❌ Gagal export item: {e}")

    # ========================
    # 📤 Export Tabel Umum (baru)
    # ========================
    @db_group.command(name="export")
    @commands.has_permissions(administrator=True)
    async def export_table(self, ctx, table: str, fmt: str = "csv"):
        """📤 Export tabel tertentu dari database (contoh: !db export npc xlsx)."""
        guild_id = ctx.guild.id

        # alias otomatis (biar !db export faction tetap jalan)
        aliases = {
            "faction": "factions",
            "favor": "favors",
            "enemy": "enemies",
            "character": "characters",
            "quest": "quests",
            "npcshop": "npc_shop",
            "hollow": "hollow_nodes",
            "visitor": "hollow_visitors",
            "event": "hollow_events",
        }
        table = aliases.get(table.lower(), table.lower())

        try:
            rows = db.fetchall(guild_id, f"SELECT * FROM {table}")
        except Exception as e:
            return await ctx.send(f"❌ Gagal membaca tabel `{table}`: {e}")

        if not rows:
            return await ctx.send(f"⚠️ Tabel `{table}` kosong atau tidak ditemukan.")

        df = pd.DataFrame(rows)
        os.makedirs("/tmp", exist_ok=True)
        filename = f"/tmp/{table}_{guild_id}.{fmt.lower()}"

        try:
            fmt = fmt.lower()
            if fmt == "csv":
                df.to_csv(filename, index=False)
            elif fmt == "xlsx":
                df.to_excel(filename, index=False)
            elif fmt == "json":
                df.to_json(filename, orient="records", indent=2)
            elif fmt == "docx":
                from docx import Document
                doc = Document()
                doc.add_heading(f"Technonesia Export – {table}", level=1)
                for _, r in df.iterrows():
                    title = str(r.get("name", f"Row {_+1}"))
                    doc.add_heading(title, level=2)
                    for k, v in r.items():
                        doc.add_paragraph(f"{k}: {v}")
                    doc.add_paragraph("—")
                doc.save(filename)
            else:
                return await ctx.send("❌ Format tidak dikenal! Gunakan: csv / xlsx / docx / json.")

            await ctx.send(
                content=f"✅ Data `{table}` diekspor sebagai **{fmt.upper()}**",
                file=discord.File(filename),
            )

        except Exception as e:
            await ctx.send(f"❌ Gagal export `{table}`: {e}")

    @db_group.command(name="finddb")
    @commands.has_permissions(administrator=True)
    async def finddb(self, ctx):
        """🔍 Cari semua file .db di seluruh project."""
        import os
        possible_dirs = [
            ".",                # root project
            "./data",           # path lama (sebelum versi baru)
            "./app/data",       # path baru (runtime container)
            "../data",          # kemungkinan parent
        ]

        found = []
        for base in possible_dirs:
            if os.path.exists(base):
                for root, dirs, files in os.walk(base):
                    for f in files:
                        if f.endswith(".db"):
                            full = os.path.join(root, f)
                            size = os.path.getsize(full) / 1024
                            found.append(f"{full} — {size:.1f} KB")

        if not found:
            return await ctx.send("📭 Tidak ada file `.db` ditemukan di direktori yang bisa diakses.")

        chunks = [found[i:i + 10] for i in range(0, len(found), 10)]
        for i, ch in enumerate(chunks, 1):
            text = "\n".join(ch)
            await ctx.send(f"📦 **Hasil Pencarian DB (bagian {i}/{len(chunks)})**\n```{text}```")

    # ========================
    # 📦 Export Semua Tabel Utama (baru)
    # ========================
    @db_group.command(name="exportall")
    @commands.has_permissions(administrator=True)
    async def export_all(self, ctx, fmt: str = "xlsx"):
        """📦 Export semua tabel utama (characters, npc, items, quests, dll)."""
        guild_id = ctx.guild.id
        fmt = fmt.lower()
        tables = [
            "characters", "npc", "items", "quests", "factions", "favors",
            "hollow_nodes", "hollow_events", "hollow_visitors", "enemies"
        ]
        os.makedirs("/tmp", exist_ok=True)

        for t in tables:
            try:
                rows = db.fetchall(guild_id, f"SELECT * FROM {t}")
                if not rows:
                    continue
                df = pd.DataFrame(rows)
                filename = f"/tmp/{t}_{guild_id}.{fmt}"
                if fmt == "xlsx":
                    df.to_excel(filename, index=False)
                else:
                    df.to_csv(filename, index=False)
                await ctx.send(f"✅ Export `{t}` selesai.", file=discord.File(filename))
            except Exception as e:
                await ctx.send(f"⚠️ Gagal export `{t}`: {e}")


async def setup(bot):
    await bot.add_cog(DbAdmin(bot))
